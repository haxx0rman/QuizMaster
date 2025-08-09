"""
Document processing module using BookWorm integration.
"""

import logging
import asyncio
from pathlib import Path
from typing import Dict, List, Any, Optional, Union
from dataclasses import dataclass

from .config import QuizMasterConfig


@dataclass
class ProcessedDocument:
    """Represents a processed document with metadata."""
    path: str
    title: str
    content: str
    metadata: Dict[str, Any]
    word_count: int
    file_size: int
    processing_time: float


class DocumentProcessor:
    """
    Document processor that wraps BookWorm functionality.
    
    This class provides a simplified interface for document processing
    that integrates well with the QuizMaster workflow.
    """
    
    def __init__(self, config: QuizMasterConfig):
        """
        Initialize the document processor.
        
        Args:
            config: QuizMasterConfig instance
        """
        self.config = config
        self.logger = logging.getLogger(__name__)
        
        # We'll initialize BookWorm components when they're needed
        # to avoid import errors if BookWorm isn't available
        self._bookworm_processor = None
        self._mindmap_generator = None
    
    @property
    def bookworm_processor(self):
        """Lazy initialization of BookWorm processor."""
        if self._bookworm_processor is None:
            try:
                from bookworm.core import DocumentProcessor as BookwormProcessor
                from bookworm.utils import BookWormConfig
                
                # Start with a basic config and set essential parameters
                bookworm_config = BookWormConfig()
                
                # Set the API keys
                if self.config.openai_api_key:
                    bookworm_config.openai_api_key = self.config.openai_api_key
                if self.config.anthropic_api_key:
                    bookworm_config.anthropic_api_key = self.config.anthropic_api_key
                if self.config.deepseek_api_key:
                    bookworm_config.deepseek_api_key = self.config.deepseek_api_key
                if self.config.gemini_api_key:
                    bookworm_config.gemini_api_key = self.config.gemini_api_key
                
                # Set other configuration
                bookworm_config.api_provider = self.config.api_provider
                bookworm_config.llm_model = self.config.llm_model
                
                # Set directory paths and ensure they exist
                working_path = Path(self.config.working_dir)
                bookworm_config.working_dir = str(working_path)
                bookworm_config.document_dir = str(working_path / "docs")
                bookworm_config.processed_dir = str(working_path / "processed_docs")
                bookworm_config.output_dir = str(working_path / "output")
                bookworm_config.log_dir = str(working_path / "logs")
                
                # Ensure all required directories exist
                working_path.mkdir(exist_ok=True)
                (working_path / "library").mkdir(exist_ok=True)  # BookWorm needs this
                (working_path / "docs").mkdir(exist_ok=True)
                (working_path / "processed_docs").mkdir(exist_ok=True)
                (working_path / "output").mkdir(exist_ok=True)
                (working_path / "logs").mkdir(exist_ok=True)
                
                self._bookworm_processor = BookwormProcessor(config=bookworm_config)
                self.logger.info("BookWorm processor initialized successfully")
                
            except ImportError as e:
                self.logger.error(f"BookWorm not available: {e}")
                raise ImportError("BookWorm package not found")
            except Exception as e:
                self.logger.error(f"Error initializing BookWorm processor: {str(e)}")
                raise RuntimeError(f"Failed to initialize BookWorm: {e}")
        return self._bookworm_processor
    
    @property
    def mindmap_generator(self):
        """Lazy initialization of mindmap generator."""
        if self._mindmap_generator is None:
            try:
                from bookworm.mindmap_generator import BookWormMindmapGenerator
                # Use the same BookWorm config as the processor
                bookworm_config = self.bookworm_processor.config
                self._mindmap_generator = BookWormMindmapGenerator(config=bookworm_config)
            except ImportError:
                self.logger.error("BookWorm mindmap generator not available.")
                raise ImportError("BookWorm package not found or incomplete")
            except Exception as e:
                self.logger.error(f"Error initializing mindmap generator: {str(e)}")
                raise RuntimeError(f"Failed to initialize mindmap generator: {e}")
        return self._mindmap_generator
    
    @property
    def knowledge_graph(self):
        """Access to the knowledge graph through BookWorm."""
        if hasattr(self.bookworm_processor, 'library_manager'):
            return self.bookworm_processor.library_manager
        else:
            # Fallback: create a simple mock for demo
            return self._create_mock_knowledge_graph()
    
    def _create_mock_knowledge_graph(self):
        """Create a mock knowledge graph for demo purposes."""
        class MockKnowledgeGraph:
            async def insert(self, content):
                # Mock insertion
                pass
            
            async def query(self, query, mode="naive"):
                # Mock query that returns sample content
                return f"Mock knowledge graph response for query: {query}"
        
        return MockKnowledgeGraph()
    
    async def process_document(self, document_path: str) -> ProcessedDocument:
        """
        Process a single document.
        
        Args:
            document_path: Path to the document to process
            
        Returns:
            ProcessedDocument object with extracted content and metadata
        """
        self.logger.info(f"Processing document: {document_path}")
        
        path = Path(document_path)
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {document_path}")
        
        # Check file size
        file_size = path.stat().st_size
        max_size = self.config.max_file_size_mb * 1024 * 1024
        if file_size > max_size:
            raise ValueError(f"File too large: {file_size} bytes (max: {max_size} bytes)")
        
        try:
            import time
            start_time = time.time()
            
            # Use BookWorm to process the document
            processed_content = await self._process_with_bookworm(str(path))
            
            processing_time = time.time() - start_time
            
            # Extract metadata
            metadata = {
                "file_type": path.suffix.lower(),
                "file_name": path.name,
                "file_path": str(path.absolute()),
                "title": processed_content.get("title", path.stem),
                "processing_method": "bookworm",
                "processed_at": time.time()
            }
            
            # Create processed document object
            processed_doc = ProcessedDocument(
                path=str(path),
                title=metadata["title"],
                content=processed_content["content"],
                metadata=metadata,
                word_count=len(processed_content["content"].split()),
                file_size=file_size,
                processing_time=processing_time
            )
            
            self.logger.info(f"Document processed successfully: {document_path} ({processed_doc.word_count} words)")
            return processed_doc
            
        except Exception as e:
            self.logger.error(f"Error processing document {document_path}: {str(e)}")
            raise
    
    async def process_multiple_documents(
        self, 
        document_paths: List[str],
        max_concurrent: Optional[int] = None
    ) -> List[ProcessedDocument]:
        """
        Process multiple documents concurrently.
        
        Args:
            document_paths: List of paths to documents
            max_concurrent: Maximum number of concurrent processing tasks
            
        Returns:
            List of ProcessedDocument objects
        """
        if max_concurrent is None:
            max_concurrent = self.config.max_concurrent_processes
        
        self.logger.info(f"Processing {len(document_paths)} documents with {max_concurrent} concurrent tasks")
        
        # Create semaphore to limit concurrency
        semaphore = asyncio.Semaphore(max_concurrent)
        
        async def process_with_semaphore(doc_path: str) -> Optional[ProcessedDocument]:
            async with semaphore:
                try:
                    return await self.process_document(doc_path)
                except Exception as e:
                    self.logger.error(f"Failed to process {doc_path}: {str(e)}")
                    return None
        
        # Process documents concurrently
        tasks = [process_with_semaphore(doc_path) for doc_path in document_paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Filter out None results and exceptions
        processed_docs = [
            doc for doc in results 
            if doc is not None and not isinstance(doc, Exception)
        ]
        
        self.logger.info(f"Successfully processed {len(processed_docs)}/{len(document_paths)} documents")
        return processed_docs
    
    async def generate_mindmap(
        self, 
        content: str, 
        title: str = "Document Analysis",
        output_format: str = "mermaid"
    ) -> Dict[str, Any]:
        """
        Generate a mindmap from document content.
        
        Args:
            content: Text content to analyze
            title: Title for the mindmap
            output_format: Output format (mermaid, html, markdown)
            
        Returns:
            Dictionary containing mindmap data and metadata
        """
        if not self.config.enable_mindmaps:
            self.logger.warning("Mindmap generation is disabled in configuration")
            return {"error": "Mindmap generation disabled"}
        
        try:
            self.logger.info(f"Generating mindmap for: {title}")
            
            # Generate mindmap using BookWorm
            mindmap_result = await self._generate_mindmap_with_bookworm(
                content, title, output_format
            )
            
            self.logger.info(f"Mindmap generated successfully for: {title}")
            return mindmap_result
            
        except Exception as e:
            self.logger.error(f"Error generating mindmap: {str(e)}")
            return {"error": str(e)}
    
    async def _process_with_bookworm(self, document_path: str) -> Dict[str, Any]:
        """
        Process document using BookWorm.
        
        This is a placeholder implementation. In practice, this would
        call the actual BookWorm document processing methods.
        """
        # For now, we'll implement basic document reading
        # This should be replaced with actual BookWorm integration
        
        try:
            path = Path(document_path)
            
            if path.suffix.lower() == '.txt':
                content = path.read_text(encoding='utf-8')
            elif path.suffix.lower() == '.md':
                content = path.read_text(encoding='utf-8')
            elif path.suffix.lower() == '.pdf':
                content = await self._extract_pdf_content(path)
            elif path.suffix.lower() in ['.docx', '.doc']:
                content = await self._extract_docx_content(path)
            else:
                # Try to read as text
                content = path.read_text(encoding='utf-8')
            
            return {
                "content": content,
                "title": path.stem,
                "source": str(path)
            }
            
        except Exception as e:
            self.logger.error(f"Error in BookWorm processing: {str(e)}")
            raise
    
    async def _extract_pdf_content(self, path: Path) -> str:
        """Extract content from PDF file."""
        try:
            if self.config.pdf_processor == "pymupdf":
                import fitz  # PyMuPDF
                doc = fitz.open(str(path))
                content = ""
                for page in doc:
                    content += page.get_text()
                doc.close()
                return content
            else:
                import pdfplumber
                content = ""
                with pdfplumber.open(str(path)) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            content += text + "\\n"
                return content
        except ImportError as e:
            self.logger.error(f"PDF processing library not available: {str(e)}")
            raise ImportError(f"PDF processor '{self.config.pdf_processor}' not available")
        except Exception as e:
            self.logger.error(f"Error extracting PDF content: {str(e)}")
            raise
    
    async def _extract_docx_content(self, path: Path) -> str:
        """Extract content from DOCX file."""
        try:
            from docx import Document
            doc = Document(str(path))
            content = ""
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\\n"
            return content
        except ImportError:
            self.logger.error("python-docx library not available")
            raise ImportError("python-docx package required for DOCX processing")
        except Exception as e:
            self.logger.error(f"Error extracting DOCX content: {str(e)}")
            raise
    
    async def _generate_mindmap_with_bookworm(
        self, 
        content: str, 
        title: str, 
        output_format: str
    ) -> Dict[str, Any]:
        """
        Generate mindmap using BookWorm.
        
        This is a placeholder implementation.
        """
        # Placeholder implementation
        return {
            "title": title,
            "format": output_format,
            "content": f"# {title}\\n\\nMindmap for content with {len(content)} characters",
            "generated_at": "2025-01-08",
            "success": True
        }
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported document formats."""
        return [
            ".txt", ".md", ".markdown",
            ".pdf",
            ".docx", ".doc",
            ".json", ".yaml", ".yml",
            ".py", ".js", ".ts", ".java", ".cpp", ".c", ".h"
        ]
    
    def is_supported_format(self, file_path: str) -> bool:
        """Check if a file format is supported."""
        path = Path(file_path)
        return path.suffix.lower() in self.get_supported_formats()
